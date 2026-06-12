"""One-shot Markdown -> .docx converter for the App Store listing.

Handles only the features used in shopify-app-store-listing.md:
- # / ## / ### headings
- Paragraphs
- > blockquotes
- - bullets / numbered lists
- | tables |
- horizontal rules ---
- **bold**, `code`, [link](href)

Run: python _md_to_docx.py shopify-app-store-listing.md shopify-app-store-listing.docx
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor


def add_runs(paragraph, text):
    """Apply inline markdown (bold / code / links) to a paragraph."""
    # Strip links to plain text + record href (Word can hyperlink but keeping simple here)
    text = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", lambda m: f"{m.group(1)} ({m.group(2)})", text)
    # Tokenise on **bold** and `code`
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    # Set base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Horizontal rule -> page-break-ish separator
        if line.strip() == "---":
            doc.add_paragraph().add_run("").add_break()
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue

        # Tables: look ahead for | header | + |---| separator
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-\s|:]+\|$", lines[i + 1]):
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
            table.style = "Light Grid Accent 1"
            for col, h in enumerate(header_cells):
                cell = table.rows[0].cells[col]
                cell.text = ""
                add_runs(cell.paragraphs[0], h)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for ridx, row in enumerate(rows):
                for col, cell_text in enumerate(row):
                    if col < len(header_cells):
                        cell = table.rows[1 + ridx].cells[col]
                        cell.text = ""
                        add_runs(cell.paragraphs[0], cell_text)
            doc.add_paragraph()
            continue

        # Blockquote
        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(lines[i][1:].lstrip())
                i += 1
            # Treat as paragraphs inside Quote style
            for ql in quote_lines:
                if not ql.strip():
                    continue
                if ql.startswith("### "):
                    p = doc.add_paragraph()
                    add_runs(p, ql[4:])
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(12)
                else:
                    p = doc.add_paragraph(style="Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else "Normal")
                    add_runs(p, ql)
            continue

        # Bullet
        if re.match(r"^\s*[-*]\s", line):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^\s*[-*]\s", "", line))
            i += 1
            continue

        # Numbered
        if re.match(r"^\s*\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\s*\d+\.\s", "", line))
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Default: paragraph
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
